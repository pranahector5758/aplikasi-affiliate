import streamlit as st
from docx import Document
from google import genai
import io

# Konfigurasi Halaman
st.set_page_config(page_title="Generator Script Affiliate", page_icon="📝", layout="centered")

st.title("Generator Script Affiliate")
st.markdown("Buat caption dan script promosi produk otomatis dalam bahasa Indonesia.")

# Form Input
product_link = st.text_input("🔗 Link Produk (Shopee, TikTok, Tokopedia, dll)", placeholder="https://...")

col1, col2, col3 = st.columns(3)
with col1:
    num_results = st.number_input("Jumlah Hasil", min_value=1, max_value=5, value=1)
with col2:
    tone = st.selectbox("Gaya Bahasa", ["Santai & Friendly", "Persuasif (Hard-selling)", "Storytelling", "Humoris", "Profesional"])
with col3:
    length = st.selectbox("Panjang Tulisan", ["Pendek (TikTok/Reels)", "Sedang (Instagram Post)", "Panjang (Blog/Facebook)"])

# Tombol Eksekusi
if st.button("✨ Buat Script Sekarang", type="primary", use_container_width=True):
    if not product_link:
        st.warning("⚠️ Silakan masukkan link produk!")
    else:
        with st.spinner("🤖 AI sedang menyusun kata-kata ajaib untukmu..."):
            try:
                # MENGAMBIL API KEY DARI BRANKAS RAHASIA STREAMLIT
                api_key = st.secrets["GEMINI_API_KEY"]
                
                # Setup AI menggunakan SDK genai terbaru
                client = genai.Client(api_key=api_key)
                
                # Prompt instruksi ke AI
                prompt = f"""
                Kamu adalah seorang copywriter affiliate marketing profesional.
                Buatkan {num_results} variasi script promosi/affiliate dalam bahasa Indonesia untuk produk pada link berikut: {product_link}.
                
                Instruksi spesifik:
                - Gaya bahasa: {tone}
                - Panjang tulisan: {length}
                - Berikan judul yang jelas untuk setiap variasi (misal: 📌 Opsi 1).
                - Tambahkan Call to Action (CTA) dan berikan placeholder [LINK PRODUK] agar mudah diganti.
                - Tambahkan hashtag yang relevan di akhir teks.
                """
                
                # Menghasilkan teks
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt
                )
                hasil_teks = response.text
                
                st.success("✅ Script berhasil dibuat!")
                
                # Tampilkan hasil
                st.markdown("### 📝 Hasil Generate")
                st.text_area("Blok semua teks di bawah atau gunakan tombol 'Copy' di pojok kanan atas kotak ini:", value=hasil_teks, height=350)
                
                # Fitur Export ke Word (.docx)
                doc = Document()
                doc.add_heading('Script Konten Affiliate', 0)
                doc.add_paragraph(hasil_teks)
                
                # Simpan ke memori sementara untuk didownload
                bio = io.BytesIO()
                doc.save(bio)
                
                st.download_button(
                    label="📄 Export Hasil ke Word (.docx)",
                    data=bio.getvalue(),
                    file_name="Script_Affiliate.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
            except KeyError:
                st.error("⚠️ Sistem belum menemukan API Key rahasia. Pastikan Anda sudah mengatur 'Secrets' di Streamlit Cloud.")
            except Exception as e:
                st.error(f"Terjadi kesalahan: {e}")